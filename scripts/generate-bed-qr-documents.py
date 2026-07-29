#!/usr/bin/env python3
"""Generate permanent bed QR codes and a print-ready Word document.

Bed numbers, crop names, and timeline targets are read directly from
js/garden-data.js via Node.js. No separate bed list is maintained here.
"""

from __future__ import annotations

import argparse
import json
import re
import subprocess
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.graphics.barcode.qrencoder import QRCode, QRErrorCorrectLevel
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


SITE_ROOT = "https://caadx317.github.io/DF-WEB/"
QR_BORDER_MODULES = 4
QR_BOX_PIXELS = 32


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
    )
    parser.add_argument("--node", default="node")
    parser.add_argument(
        "--homepage-only",
        action="store_true",
        help="Generate only Homepage_QR.png and Homepage_QR.pdf.",
    )
    return parser.parse_args()


def load_beds(project_root: Path, node_executable: str) -> list[dict]:
    data_path = project_root / "js" / "garden-data.js"
    node_script = r"""
const fs = require("fs");
const vm = require("vm");
const source = fs.readFileSync(process.argv[1], "utf8");
const context = {};
vm.createContext(context);
vm.runInContext(`${source}\nthis.__gardenData = gardenData;`, context);
const beds = context.__gardenData.beds
  .filter((bed) => /^bed-\d+$/.test(bed.id))
  .map((bed) => ({
    number: Number(bed.id.slice(4)),
    id: bed.id,
    cropRaw: String(bed.crop || ""),
    cropName: String(bed.crop || "").replace(/_/g, " "),
    timelineLink: bed.link
  }))
  .sort((a, b) => a.number - b.number);
process.stdout.write(JSON.stringify(beds));
"""
    completed = subprocess.run(
        [node_executable, "-e", node_script, str(data_path)],
        check=True,
        text=True,
        capture_output=True,
    )
    beds = json.loads(completed.stdout)
    if not beds:
        raise RuntimeError("No numbered beds were found in js/garden-data.js")
    numbers = [bed["number"] for bed in beds]
    if len(numbers) != len(set(numbers)):
        raise RuntimeError("Duplicate numbered beds were found in js/garden-data.js")
    for bed in beds:
        if not bed["cropName"].strip():
            raise RuntimeError(f'{bed["id"]} has no crop name')
        bed["url"] = f'{SITE_ROOT}#bed-{bed["number"]}'
    return beds


def safe_crop_filename(crop_name: str) -> str:
    name = re.sub(r"\s+", "_", crop_name.strip())
    name = re.sub(r"[^A-Za-z0-9_-]", "", name)
    name = re.sub(r"_+", "_", name).strip("_.")
    return name or "Unnamed_Crop"


def build_qr_matrix(url: str) -> list[list[bool]]:
    qr = QRCode(None, QRErrorCorrectLevel.H)
    qr.addData(url)
    qr.make()
    return [[bool(cell) for cell in row] for row in qr.modules]


def render_qr_png(
    matrix: list[list[bool]],
    output_path: Path,
    *,
    border_modules: int = QR_BORDER_MODULES,
    box_pixels: int = QR_BOX_PIXELS,
) -> None:
    module_count = len(matrix)
    canvas_modules = module_count + (2 * border_modules)
    size = canvas_modules * box_pixels
    image = Image.new("RGB", (size, size), "white")
    pixels = image.load()

    for row, values in enumerate(matrix):
        for col, is_dark in enumerate(values):
            if not is_dark:
                continue
            x0 = (col + border_modules) * box_pixels
            y0 = (row + border_modules) * box_pixels
            for y in range(y0, y0 + box_pixels):
                for x in range(x0, x0 + box_pixels):
                    pixels[x, y] = (0, 0, 0)

    image.save(output_path, format="PNG", optimize=True, dpi=(300, 300))


def validate_qr_bitmap(
    matrix: list[list[bool]],
    output_path: Path,
    expected_url: str,
    *,
    border_modules: int = QR_BORDER_MODULES,
    box_pixels: int = QR_BOX_PIXELS,
) -> None:
    with Image.open(output_path) as image:
        image = image.convert("RGB")
        module_count = len(matrix)
        expected_size = (module_count + 2 * border_modules) * box_pixels
        if image.size != (expected_size, expected_size):
            raise RuntimeError(f"Unexpected QR dimensions for {output_path.name}")
        if image.getpixel((0, 0)) != (255, 255, 255):
            raise RuntimeError(f"Missing white quiet zone in {output_path.name}")
        for row, values in enumerate(matrix):
            for col, is_dark in enumerate(values):
                x = (col + border_modules) * box_pixels + box_pixels // 2
                y = (row + border_modules) * box_pixels + box_pixels // 2
                expected = (0, 0, 0) if is_dark else (255, 255, 255)
                if image.getpixel((x, y)) != expected:
                    raise RuntimeError(f"QR module mismatch in {output_path.name}")
    if not expected_url.startswith(SITE_ROOT):
        raise RuntimeError(f"Unexpected URL encoded for {output_path.name}")


def set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    color: str = "000000",
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 30, "000000", 0, 6),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_bed_page(doc: Document, bed: dict, qr_path: Path) -> None:
    bed_heading = doc.add_paragraph()
    bed_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bed_heading.paragraph_format.space_before = Pt(0)
    bed_heading.paragraph_format.space_after = Pt(6)
    bed_heading.paragraph_format.keep_with_next = True
    set_run_font(
        bed_heading.add_run(f'Bed {bed["number"]}'),
        size=30,
        bold=True,
    )

    crop_heading = doc.add_paragraph()
    crop_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crop_heading.paragraph_format.space_before = Pt(0)
    crop_heading.paragraph_format.space_after = Pt(22)
    crop_heading.paragraph_format.keep_with_next = True
    set_run_font(
        crop_heading.add_run(bed["cropName"]),
        size=24,
        bold=True,
    )

    qr_paragraph = doc.add_paragraph()
    qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_paragraph.paragraph_format.space_before = Pt(0)
    qr_paragraph.paragraph_format.space_after = Pt(18)
    qr_paragraph.paragraph_format.keep_with_next = True
    qr_run = qr_paragraph.add_run()
    qr_run.add_picture(str(qr_path), width=Inches(3.75), height=Inches(3.75))

    url_paragraph = doc.add_paragraph()
    url_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_paragraph.paragraph_format.space_before = Pt(0)
    url_paragraph.paragraph_format.space_after = Pt(0)
    url_paragraph.paragraph_format.keep_together = True
    set_run_font(url_paragraph.add_run(bed["url"]), size=9, color="333333")


def build_document(beds: list[dict], project_root: Path) -> Path:
    output_path = project_root / "Oak_Creek_Bed_QR_Codes.docx"
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Oak Creek Bed QR Codes"
    doc.core_properties.subject = "Permanent garden bed links"

    for index, bed in enumerate(beds):
        add_bed_page(doc, bed, Path(bed["qrPath"]))
        if index < len(beds) - 1:
            doc.add_page_break()

    doc.save(output_path)
    return output_path


def build_homepage_pdf(qr_path: Path, output_path: Path) -> None:
    page_pixel_width = 3400
    page_pixel_height = 4400
    page_image = Image.new(
        "RGB",
        (page_pixel_width, page_pixel_height),
        "white",
    )
    draw = ImageDraw.Draw(page_image)

    bold_font_paths = (
        Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf"),
        Path("/Library/Fonts/Arial Bold.ttf"),
    )
    regular_font_paths = (
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/Library/Fonts/Arial.ttf"),
    )
    bold_font_path = next((path for path in bold_font_paths if path.is_file()), None)
    regular_font_path = next(
        (path for path in regular_font_paths if path.is_file()),
        None,
    )
    if not bold_font_path or not regular_font_path:
        raise RuntimeError("Arial fonts required for the homepage PDF were not found")

    title_font = ImageFont.truetype(str(bold_font_path), 140)
    url_font = ImageFont.truetype(str(regular_font_path), 60)
    title = "Oak Creek Dry Farming Greens"

    title_box = draw.textbbox((0, 0), title, font=title_font)
    title_x = (page_pixel_width - (title_box[2] - title_box[0])) // 2
    draw.text((title_x, 300), title, fill="black", font=title_font)

    with Image.open(qr_path) as qr_source:
        qr_image = qr_source.convert("RGB").resize(
            (1960, 1960),
            Image.Resampling.NEAREST,
        )
    qr_x = (page_pixel_width - qr_image.width) // 2
    qr_y = 920
    page_image.paste(qr_image, (qr_x, qr_y))

    url_box = draw.textbbox((0, 0), SITE_ROOT, font=url_font)
    url_x = (page_pixel_width - (url_box[2] - url_box[0])) // 2
    draw.text((url_x, 3100), SITE_ROOT, fill="black", font=url_font)

    page_width, page_height = letter
    with tempfile.TemporaryDirectory() as temp_dir:
        page_image_path = Path(temp_dir) / "homepage-qr-page.png"
        page_image.save(page_image_path, format="PNG", optimize=True, dpi=(400, 400))

        pdf = canvas.Canvas(str(output_path), pagesize=letter)
        pdf.setTitle("Oak Creek Dry Farming Greens Homepage QR Code")
        pdf.drawImage(
            str(page_image_path),
            0,
            0,
            width=page_width,
            height=page_height,
        )
        pdf.showPage()
        pdf.save()


def generate_homepage_assets(project_root: Path) -> dict[str, str | int]:
    qr_dir = project_root / "generated_qr_codes"
    qr_dir.mkdir(exist_ok=True)
    qr_path = qr_dir / "Homepage_QR.png"
    pdf_path = project_root / "Homepage_QR.pdf"
    border_modules = 8
    box_pixels = 48

    matrix = build_qr_matrix(SITE_ROOT)
    render_qr_png(
        matrix,
        qr_path,
        border_modules=border_modules,
        box_pixels=box_pixels,
    )
    validate_qr_bitmap(
        matrix,
        qr_path,
        SITE_ROOT,
        border_modules=border_modules,
        box_pixels=box_pixels,
    )
    build_homepage_pdf(qr_path, pdf_path)

    with Image.open(qr_path) as image:
        width, height = image.size
    return {
        "url": SITE_ROOT,
        "png": str(qr_path),
        "pdf": str(pdf_path),
        "width": width,
        "height": height,
        "quietZoneModules": border_modules,
    }


def main() -> None:
    args = parse_args()
    project_root = args.project_root.resolve()
    qr_dir = project_root / "generated_qr_codes"
    qr_dir.mkdir(exist_ok=True)

    if args.homepage_only:
        print(json.dumps(generate_homepage_assets(project_root), indent=2))
        return

    beds = load_beds(project_root, args.node)
    for bed in beds:
        filename = (
            f'Bed_{bed["number"]:02d}_{safe_crop_filename(bed["cropName"])}_QR.png'
        )
        qr_path = qr_dir / filename
        matrix = build_qr_matrix(bed["url"])
        render_qr_png(matrix, qr_path)
        validate_qr_bitmap(matrix, qr_path, bed["url"])
        bed["qrPath"] = str(qr_path)
        bed["qrFilename"] = filename

    manifest_path = qr_dir / "bed_qr_manifest.json"
    manifest_path.write_text(
        json.dumps(beds, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    docx_path = build_document(beds, project_root)
    print(
        json.dumps(
            {
                "beds": len(beds),
                "qrCodes": len(beds),
                "manifest": str(manifest_path),
                "docx": str(docx_path),
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
